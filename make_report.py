"""숙제 보고서 DOCX 생성"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 기본 스타일 설정
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5

# 여백 설정
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_title(text, level=0):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "맑은 고딕"
    return h


def add_body(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run.font.size = Pt(10.5)
    run.bold = bold
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = "맑은 고딕"
                run.font.size = Pt(10)
    # 데이터
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(10)
    doc.add_paragraph()  # 표 뒤 간격
    return table


def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


# ─────────────────────────────────────
# 표지
# ─────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading("Munger Lattice Stock Analyzer", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("찰리 멍거의 격자모델 기반 AI 멀티에이전트 주식 분석 시스템")
run.font.size = Pt(14)
run.font.name = "맑은 고딕"
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

add_table(
    ["항목", "내용"],
    [
        ["이름", ""],
        ["학번", ""],
        ["과정명", ""],
        ["제출일", "2026년 3월 27일"],
    ],
)

doc.add_page_break()

# ─────────────────────────────────────
# 1. 설명
# ─────────────────────────────────────
add_title("1. 설명 (Description)", level=1)

add_body(
    "이 AI는 찰리 멍거(Charlie Munger)의 격자모델(Mental Model Lattice) "
    "의사결정 프레임워크를 주식 투자 분석에 적용한 멀티에이전트 시스템이다."
)
add_body(
    "사용자가 종목명을 입력하면, 물리학/생물학/경제학/심리학/역발상의 "
    "5가지 학문적 관점에서 각각 독립적으로 분석한 후, "
    "수석 분석관이 이를 종합하여 최종 투자 보고서를 생성한다."
)

add_body("핵심 차별점:", bold=True)
bullets = [
    "단일 관점이 아닌 다학제적 분석 (멍거의 핵심 철학 반영)",
    "5개의 독립 AI 에이전트가 병렬로 동시 분석 (속도 최적화)",
    "역발상 에이전트가 다른 에이전트들의 결론에 반론을 제기 (편향 방지)",
    "yfinance 재무 데이터 + Google News RSS 뉴스를 실시간 자동 수집",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

# ─────────────────────────────────────
# 2. 시스템 아키텍처
# ─────────────────────────────────────
add_title("2. 시스템 아키텍처", level=1)

add_code_block(
    "[사용자] → 종목명 입력 → [Flask 웹앱 (EC2)]\n"
    "                              ↓\n"
    "                     [데이터 수집 모듈]\n"
    "                     ├── yfinance → 재무지표\n"
    "                     └── Google News RSS → 최신 뉴스 10건\n"
    "                              ↓\n"
    "         ┌──── 5개 AI 에이전트 (병렬 실행) ────┐\n"
    "         │  ⚛️ 물리학: 안전마진, 평균회귀       │\n"
    "         │  🧬 생물학: 임계점, 진화/적응        │\n"
    "         │  📊 경제학: 금리, 투자자수익률       │\n"
    "         │  🧠 심리학: 인지편향, 인센티브       │\n"
    "         │  🔄 역발상: 실패시나리오, 반론       │\n"
    "         │  (Claude Haiku 4.5 × 5건 병렬)     │\n"
    "         └──────────────┬──────────────────────┘\n"
    "                        ↓\n"
    "         ┌──── 수석 분석관 (취합) ─────────────┐\n"
    "         │  5개 분석 결과 종합                  │\n"
    "         │  종합 점수 + 투자의견 + 보고서 생성   │\n"
    "         │  (Claude Sonnet 4.5 × 1건)         │\n"
    "         └────────────────────────────────────┘\n"
    "                        ↓\n"
    "         [웹 페이지에 탭 형태로 보고서 출력]"
)

# ─────────────────────────────────────
# 3. 에이전트별 시스템 프롬프트
# ─────────────────────────────────────
add_title("3. 에이전트별 시스템 프롬프트", level=1)

agents = [
    (
        "3-1. 물리학 분석관 (⚛️)",
        "평균회귀와 안전마진 관점에서 밸류에이션 분석",
        "당신은 찰리 멍거의 격자모델 중 '물리학' 관점에서 주식을 분석하는 전문가입니다.\n\n"
        "다음 원칙을 적용하여 분석하세요:\n\n"
        "1. 평균회귀와 안전마진\n"
        "- 역사적 평균(PER, PBR, 주가 범위) 대비 현재 수준?\n"
        "- 52주 최고가, 최저가 대비 현재 수준\n"
        "- 경쟁사 대비 밸류에이션 수준(해당섹터의 경쟁기업의 지표 직접비교표 추가)\n\n"
        "제공된 재무데이터와 뉴스를 근거로 구체적 수치를 인용하며 분석하세요.\n"
        "마지막에 이 관점에서의 점수(100점 만점)와 한줄 결론을 제시하세요.",
    ),
    (
        "3-2. 생물학 분석관 (🧬)",
        "임계점과 진화/적응 관점에서 기업의 성장 단계 및 생존력 분석",
        "당신은 찰리 멍거의 격자모델 중 '생물학' 관점에서 주식을 분석하는 전문가입니다.\n\n"
        "다음 두 가지 원칙을 적용하여 분석하세요:\n\n"
        "1. 임계점 (Tipping Point)\n"
        "- 이 기업/산업이 폭발적 성장의 임계점에 있는가?\n"
        "- 임계질량을 달성했는가, 아니면 아직 도달 전인가?\n\n"
        "2. 진화와 적응\n"
        "- 경쟁자 대비 생존 능력(적응력, 혁신성)은 어떠한가?\n"
        "- 도태 위험이 있는가? (기술 변화, 규제 변화, 소비자 변화)\n\n"
        "제공된 재무데이터와 뉴스를 근거로 구체적 수치를 인용하며 분석하세요.\n"
        "마지막에 이 관점에서의 점수(100점 만점)와 한줄 결론을 제시하세요.",
    ),
    (
        "3-3. 경제학 분석관 (📊)",
        "거시경제 환경, 투자자 수익률, 비즈니스 구조 분석",
        "당신은 찰리 멍거의 격자모델 중 '경제학' 관점에서 주식을 분석하는 전문가입니다.\n\n"
        "다음 원칙을 적용하여 분석하세요:\n\n"
        "1. 금리\n"
        "- 금리 변동이 이 기업의 자본비용, 할인율, 차입구조에 미치는 영향은?\n"
        "- 현재 금리 환경이 이 종목에 유리한가 불리한가?\n\n"
        "2. 투자자 수익률\n"
        "- 단순 PER 영업이익률 기준 분석이 아니라, 투자자에게 돌아오는 배당+자사주소각을\n"
        "  투자자수익률로 환산하여 이를 핵심지표로 채권 수익률과 비교.\n\n"
        "3. 인플레이션\n"
        "- 이 기업은 인플레이션을 가격에 전가할 수 있는 가격결정력이 있는가?\n\n"
        "4. 비즈니스 구조\n"
        "- 규모의 경제가 작동하는가? 경쟁우위의 원천은?\n"
        "- 현금흐름의 질과 지속가능성은?\n\n"
        "제공된 재무데이터와 뉴스를 근거로 구체적 수치를 인용하며 분석하세요.\n"
        "마지막에 이 관점에서의 점수(100점 만점)와 한줄 결론을 제시하세요.",
    ),
    (
        "3-4. 심리학 분석관 (🧠)",
        "투자자 인지편향 점검 및 경영진 인센티브 분석",
        "당신은 찰리 멍거의 격자모델 중 '심리학' 관점에서 주식을 분석하는 전문가입니다.\n\n"
        "다음 편향들과 심리 요인을 점검하세요:\n\n"
        "1. 군중심리 — 과도한 낙관 또는 비관에 빠져 있지 않은가?\n"
        "2. 호감/비호감 편향 — 브랜드 호감도가 투자 판단을 왜곡하고 있지 않은가?\n"
        "3. 이념편향 — 특정 테마(AI, ESG 등)에 대한 맹신이 영향을 주고 있지 않은가?\n"
        "4. 고정관념 편향 — 업종에 대한 고정관념이 기회나 위험을 가리고 있지 않은가?\n"
        "5. 우위이지만 열위 선택 — 더 나은 대안이 있는데 익숙함 때문에 선택하는 것은 아닌가?\n"
        "6. 인간관계/인센티브 — 경영진의 인센티브 구조가 주주 이익과 일치하는가?\n\n"
        "제공된 재무데이터와 뉴스를 근거로 구체적 사례를 들며 분석하세요.\n"
        "마지막에 이 관점에서의 점수(100점 만점)와 한줄 결론을 제시하세요.",
    ),
    (
        "3-5. 역발상 분석관 (🔄)",
        "실패 시나리오 도출 및 반론 제기 (Devil's Advocate)",
        "당신은 찰리 멍거의 격자모델 중 '뒤집어 생각하라(Inversion)' 원칙을 전담하는 분석관입니다.\n\n"
        '멍거는 말했습니다: "문제를 뒤집어라. 항상 뒤집어라."\n\n'
        "1. 실패 시나리오\n"
        "- 이 투자가 실패하려면 어떤 조건이 필요한가?\n"
        "- 그 조건이 실현될 확률은 얼마나 되는가?\n\n"
        "2. 반론 (Devil's Advocate)\n"
        "- 이 종목에 대한 긍정론의 가장 큰 약점은?\n"
        "- 시장이 간과하고 있는 숨겨진 리스크는?\n\n"
        "냉정하고 비판적으로 분석하세요.\n"
        "마지막에 위험도 점수(100점 만점, 높을수록 위험)와 한줄 결론을 제시하세요.",
    ),
]

for title_text, role_text, prompt_text in agents:
    add_title(title_text, level=2)
    add_body(f"역할: {role_text}", bold=True)
    add_body("시스템 프롬프트:", bold=True)
    add_code_block(prompt_text)

add_title("3-6. 수석 분석관 (📋) — 취합", level=2)
add_body(
    "5개 에이전트의 분석 결과를 입력받아, 종합 점수/투자의견/핵심 강점 3가지/"
    "핵심 리스크 3가지/최종 종합 의견을 포함한 구조화된 보고서를 출력한다. "
    "단순 평균이 아닌 가중 판단을 하며, 분석 간 모순점이나 시너지를 찾아내어 "
    "더 깊은 통찰을 제공하도록 설계하였다."
)

doc.add_page_break()

# ─────────────────────────────────────
# 4. 제작 의도
# ─────────────────────────────────────
add_title("4. 제작 의도 (Development Purpose)", level=1)

add_title("4.1 타겟 사용자", level=2)
for t in [
    "개인 투자자 (특히 가치투자에 관심 있는 투자자)",
    "투자 공부를 시작하는 학생",
    "다양한 관점에서 종목을 검토하고 싶은 투자자",
]:
    doc.add_paragraph(t, style="List Bullet")

add_title("4.2 목표 가치", level=2)
add_table(
    ["가치", "설명"],
    [
        ["다학제적 분석", "물리/생물/경제/심리/역발상 5가지 관점에서 종합 판단"],
        ["편향 방지", "역발상 에이전트가 긍정론에 반론을 제기하여 확증편향 차단"],
        ["시간 절약", "30초~1분 안에 5개 관점의 심층 분석 보고서 자동 생성"],
        ["투자 교육", "멍거의 격자모델을 실제 종목에 적용하여 투자 프레임워크 학습"],
    ],
)

add_title("4.3 성과 측정 방법 (KPI)", level=2)
add_table(
    ["지표", "측정 방법"],
    [
        ["분석 완료율", "종목 입력 대비 보고서 정상 생성 비율 (목표: 95% 이상)"],
        ["응답 시간", "분석 요청 → 보고서 출력 소요 시간 (목표: 60초 이내)"],
        ["비용 효율", "보고서 1건당 API 비용 (현재: ~$0.18)"],
        ["분석 품질", "에이전트 간 관점 차이 발생 비율 (다양성 확보)"],
    ],
)

add_title("4.4 프롬프트 설계 의도", level=2)
add_table(
    ["설계 원칙", "적용 방법"],
    [
        ["역할 특화", "각 에이전트에 명확한 학문적 관점과 분석 프레임워크 부여"],
        ["수치 근거 강제", '"구체적 수치를 인용하며 분석하세요" — 모호한 정성적 판단 방지'],
        ["점수 체계", "100점 만점 점수 + 한줄 결론 → 정량적 비교 가능"],
        ["역발상 독립성", "점수 방향을 반전시켜(높을수록 위험) 독립적 관점 유지"],
        ["투자자 수익률", "배당+자사주소각 기반 실질 수익률을 채권 수익률과 비교"],
    ],
)

add_title("4.5 도메인 지식 자료 주입", level=2)
add_table(
    ["자료", "용도"],
    [
        ["yfinance API", "실시간 재무 데이터 자동 수집 (PER, PBR, ROE, 매출 등)"],
        ["Google News RSS", "한국어/영어 최신 뉴스 10건 자동 수집"],
        ["한국 종목 매핑 테이블", '"삼성전자" → "005930.KS" 등 30개 주요 종목 자동 변환'],
    ],
)

add_body(
    "별도 파일 업로드 없이, 실시간 데이터 수집 모듈을 통해 "
    "항상 최신 데이터를 에이전트에 주입한다."
)

doc.add_page_break()

# ─────────────────────────────────────
# 5. 제작 결과보고
# ─────────────────────────────────────
add_title("5. 제작 결과보고", level=1)

add_title("5.1 시스템 구성", level=2)
add_table(
    ["항목", "내용"],
    [
        ["서비스 주소", "http://ec2-13-61-13-72.eu-north-1.compute.amazonaws.com"],
        ["소스 코드", "https://github.com/hwayobi2020/stock_analyst"],
        ["백엔드", "Python 3.9 + Flask"],
        ["AI 모델", "Claude Haiku 4.5 (에이전트) + Claude Sonnet 4.5 (수석)"],
        ["데이터", "yfinance (재무) + Google News RSS (뉴스)"],
        ["인프라", "AWS EC2 (Amazon Linux 2023, eu-north-1)"],
        ["비용", "분석 1건당 약 $0.18 (약 240원)"],
    ],
)

add_title("5.2 시작화면", level=2)
add_body("[아래에 시작화면 스크린샷을 삽입하세요]")
doc.add_paragraph()

add_title("5.3 예상 질문과 답변", level=2)
add_body("[아래에 실제 분석 결과 스크린샷을 삽입하세요]")
doc.add_paragraph()
add_body('입력: "삼성전자"', bold=True)
doc.add_paragraph()
add_body("종합 보고서 출력 (일부):", bold=True)
add_body(
    "종합 점수: 58점 / 100점\n"
    "투자의견: 중립 (Hold)\n\n"
    "핵심 요약:\n"
    "1. 저평가는 맞지만 저평가에는 이유가 있다 — PER 6.26은 메모리 반도체 사이클 저점과 "
    "구조적 성장성 의문이 동시에 반영된 수치\n"
    "2. 현금창출 능력은 우수하나 진화 실패 위험 감지 — 영업이익률 21.3%의 수익성은 건재하지만 "
    "AI 반도체 경쟁에서의 구조적 뒤처짐이 치명적\n"
    "3. 안전마진은 존재하나 촉매 부재 — 부채비율 5.79%의 재무 건전성과 대규모 현금흐름은 "
    "하방을 방어하지만, 재평가 계기가 불명확"
)

doc.add_page_break()

# ─────────────────────────────────────
# 6. 개선 방향
# ─────────────────────────────────────
add_title("6. 개선 방향", level=1)
add_table(
    ["항목", "현재", "개선안"],
    [
        ["뉴스 분석", "제목만 수집", "기사 본문 요약까지 포함"],
        ["비용 절감", "건당 $0.18", "캐싱 적용 시 90% 절감 가능"],
        ["데이터", "yfinance 기본 지표", "분기별 실적, 컨센서스, 목표주가 추가"],
        ["시각화", "텍스트 보고서", "주가 차트, 재무지표 그래프 추가"],
        ["사용자 경험", "단순 입력/출력", "분석 히스토리 저장, 종목 비교 기능"],
    ],
)

# 면책사항
doc.add_paragraph()
add_body("면책사항", bold=True)
add_body(
    "본 시스템은 찰리 멍거의 격자모델을 기반으로 AI가 생성한 분석 도구로, "
    "투자 권유가 아닙니다. 투자 결정은 반드시 본인의 판단과 책임 하에 이루어져야 합니다."
)

# 저장
doc.save("D:/projects/stock_analyst/report.docx")
print("report.docx 생성 완료")
